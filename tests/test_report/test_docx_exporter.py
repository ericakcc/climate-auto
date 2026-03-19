"""Tests for DOCX report exporter (Markdown-based)."""

from pathlib import Path

from docx import Document

from climate_auto.report.docx_exporter import generate_docx_from_markdown

_SAMPLE_MD = """\
# TACOCO 天氣討論會 - 2026-03-19

> 自動產生於 2026-03-19 12:06
>
> 報告圖片: 2 張 | 下載成功: 50 張

---

## 1. 當日回顧

### I. 分析場

![500hPa analysis](1_review/analysis/ECMWF500.gif)

Strong trough over northeast Asia.

![850hPa wind analysis](1_review/analysis/ECMWF850.gif)

Northeast monsoon dominates.

---

## 缺少的資料

- `4_context/mjo/mjo_rmm.phase.Last40days.gif`
"""

_SAMPLE_MD_NO_ANALYSIS = """\
# TACOCO 天氣討論會 - 2026-03-19

> 自動產生於 2026-03-19 12:06

---

## 1. 當日回顧

### I. 分析場

![500hPa analysis](1_review/analysis/ECMWF500.gif)

*待分析*

![850hPa wind analysis](1_review/analysis/ECMWF850.gif)

*待分析*
"""


def _create_fake_image(report_dir: Path, relative_path: str) -> None:
    """Create a minimal valid GIF file for testing."""
    image_path = report_dir / relative_path
    image_path.parent.mkdir(parents=True, exist_ok=True)
    gif_bytes = (
        b"GIF89a\x01\x00\x01\x00\x80\x00\x00"
        b"\xff\xff\xff\x00\x00\x00"
        b"!\xf9\x04\x00\x00\x00\x00\x00"
        b",\x00\x00\x00\x00\x01\x00\x01\x00\x00\x02\x02D\x01\x00;"
    )
    image_path.write_bytes(gif_bytes)


def _write_md(report_dir: Path, content: str) -> Path:
    md_path = report_dir / "daily_report.md"
    md_path.write_text(content, encoding="utf-8")
    return md_path


class TestGenerateDocxFromMarkdown:
    def test_creates_docx_file(self, tmp_path: Path) -> None:
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF500.gif")
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF850.gif")
        md_path = _write_md(tmp_path, _SAMPLE_MD)

        result = generate_docx_from_markdown(md_path, tmp_path)

        assert result.exists()
        assert result.name == "daily_report.docx"

    def test_document_structure(self, tmp_path: Path) -> None:
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF500.gif")
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF850.gif")
        md_path = _write_md(tmp_path, _SAMPLE_MD)

        result = generate_docx_from_markdown(md_path, tmp_path)
        doc = Document(str(result))

        all_texts = [p.text for p in doc.paragraphs]
        assert any("TACOCO" in t and "2026-03-19" in t for t in all_texts)
        assert any("當日回顧" in t for t in all_texts)
        assert any("分析場" in t for t in all_texts)

    def test_contains_analysis_text(self, tmp_path: Path) -> None:
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF500.gif")
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF850.gif")
        md_path = _write_md(tmp_path, _SAMPLE_MD)

        result = generate_docx_from_markdown(md_path, tmp_path)
        doc = Document(str(result))

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Strong trough" in all_text
        assert "Northeast monsoon" in all_text

    def test_skips_pending_analysis(self, tmp_path: Path) -> None:
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF500.gif")
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF850.gif")
        md_path = _write_md(tmp_path, _SAMPLE_MD_NO_ANALYSIS)

        result = generate_docx_from_markdown(md_path, tmp_path)
        doc = Document(str(result))

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "待分析" not in all_text

    def test_missing_image_shows_placeholder(self, tmp_path: Path) -> None:
        # Don't create images
        md_path = _write_md(tmp_path, _SAMPLE_MD)

        result = generate_docx_from_markdown(md_path, tmp_path)
        doc = Document(str(result))

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "圖片缺失" in all_text

    def test_missing_patterns_section(self, tmp_path: Path) -> None:
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF500.gif")
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF850.gif")
        md_path = _write_md(tmp_path, _SAMPLE_MD)

        result = generate_docx_from_markdown(md_path, tmp_path)
        doc = Document(str(result))

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "mjo_rmm" in all_text

    def test_images_are_embedded(self, tmp_path: Path) -> None:
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF500.gif")
        _create_fake_image(tmp_path, "1_review/analysis/ECMWF850.gif")
        md_path = _write_md(tmp_path, _SAMPLE_MD)

        result = generate_docx_from_markdown(md_path, tmp_path)
        doc = Document(str(result))

        assert len(doc.inline_shapes) == 2

    def test_metadata_is_included(self, tmp_path: Path) -> None:
        md_path = _write_md(tmp_path, _SAMPLE_MD)

        result = generate_docx_from_markdown(md_path, tmp_path)
        doc = Document(str(result))

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "自動產生於" in all_text
        assert "報告圖片" in all_text
