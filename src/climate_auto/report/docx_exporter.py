"""Convert a Markdown daily report to Word (.docx) format."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from loguru import logger

# Regex patterns for Markdown elements
_RE_H1 = re.compile(r"^## \d+\.\s+(.+)$")
_RE_H2 = re.compile(r"^### [IVXLCDM]+\.\s+(.+)$")
_RE_IMAGE = re.compile(r"^!\[(.+?)\]\((.+?)\)$")
_RE_TITLE = re.compile(r"^# (.+)$")
_RE_BLOCKQUOTE = re.compile(r"^>\s*(.*)$")
_RE_MISSING_HEADER = re.compile(r"^## 缺少的資料$")
_RE_BULLET = re.compile(r"^- (.+)$")


def generate_docx_from_markdown(md_path: Path, report_dir: Path) -> Path:
    """Generate a .docx report by parsing an existing Markdown report.

    Args:
        md_path: Path to the daily_report.md file.
        report_dir: Directory containing report images.

    Returns:
        Path to the generated .docx file.
    """
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Title: # TACOCO ...
        if m := _RE_TITLE.match(line):
            doc.add_heading(m.group(1), level=0)
            i += 1
            continue

        # Blockquote metadata
        if m := _RE_BLOCKQUOTE.match(line):
            text = m.group(1)
            if text:
                meta = doc.add_paragraph()
                meta.style = doc.styles["No Spacing"]
                run = meta.add_run(text)
                run.font.size = Pt(9)
                run.font.italic = True
            i += 1
            continue

        # Horizontal rule
        if line == "---":
            i += 1
            continue

        # Missing data header
        if _RE_MISSING_HEADER.match(line):
            doc.add_heading("缺少的資料", level=2)
            i += 1
            continue

        # Section heading: ## 1. xxx
        if m := _RE_H1.match(line):
            doc.add_heading(line[3:], level=1)  # Keep "1. xxx" format
            i += 1
            continue

        # Subsection heading: ### I. xxx
        if m := _RE_H2.match(line):
            doc.add_heading(line[4:], level=2)  # Keep "I. xxx" format
            i += 1
            continue

        # Image: ![desc](path)
        if m := _RE_IMAGE.match(line):
            desc, rel_path = m.group(1), m.group(2)

            # Image description (bold)
            desc_para = doc.add_paragraph()
            desc_run = desc_para.add_run(desc)
            desc_run.bold = True

            # Insert image
            image_path = report_dir / rel_path
            if image_path.exists():
                try:
                    doc.add_picture(str(image_path), width=Cm(15))
                except Exception:
                    logger.warning("Failed to embed image: {}", rel_path)
                    err_para = doc.add_paragraph()
                    err_run = err_para.add_run(f"[圖片無法嵌入: {rel_path}]")
                    err_run.font.italic = True
            else:
                missing_para = doc.add_paragraph()
                missing_run = missing_para.add_run(f"[圖片缺失: {rel_path}]")
                missing_run.font.italic = True

            # Collect analysis text (lines after image until next structure)
            i += 1
            analysis_lines: list[str] = []
            while i < len(lines):
                next_line = lines[i].rstrip()
                if (
                    not next_line
                    or _RE_H1.match(next_line)
                    or _RE_H2.match(next_line)
                    or _RE_IMAGE.match(next_line)
                    or _RE_MISSING_HEADER.match(next_line)
                    or next_line == "---"
                ):
                    if not next_line:
                        i += 1
                        continue
                    break
                analysis_lines.append(next_line)
                i += 1

            analysis_text = "\n".join(analysis_lines).strip()
            if analysis_text and analysis_text != "*待分析*":
                doc.add_paragraph(analysis_text)
            continue

        # Bullet list item
        if m := _RE_BULLET.match(line):
            doc.add_paragraph(m.group(1), style="List Bullet")
            i += 1
            continue

        # Skip empty or unrecognized lines
        i += 1

    output_path = report_dir / "daily_report.docx"
    doc.save(str(output_path))
    logger.info("DOCX report generated: {}", output_path)

    return output_path
